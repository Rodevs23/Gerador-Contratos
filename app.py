import streamlit as st
from docx import Document
from pathlib import Path
import json
import io
import os
import logging
from datetime import datetime
import re

# Configuração do logging
logging.basicConfig(
    filename='app.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

class GeradorContratosStreamlit:
    def __init__(self):
        st.set_page_config(
            page_title="Gerador de Contratos",
            layout="wide",
            page_icon="📄"
        )
        self.initialize_session_state()
        self.setup_folders()

    def setup_folders(self):
        """Cria as pastas necessárias para o funcionamento do sistema"""
        for folder in ['templates', 'backups', 'temp']:
            Path(folder).mkdir(exist_ok=True)

    def initialize_session_state(self):
        """Inicializa as variáveis de estado do Streamlit"""
        if 'current_content' not in st.session_state:
            st.session_state.current_content = ""
        if 'variables' not in st.session_state:
            st.session_state.variables = {}
        if 'original_file' not in st.session_state:
            st.session_state.original_file = None
        # Guardamos o conteúdo binário do doc em memória, para editar quando necessário
        if 'doc_data' not in st.session_state:
            st.session_state.doc_data = None

    def load_document(self, file):
        """
        Carrega o documento (sem cache para evitar problemas de atualização).
        Pode receber um arquivo de upload (BytesIO) ou um caminho (Path).
        """
        try:
            if isinstance(file, (str, Path)):
                return Document(file)
            else:
                # Se for binário (BytesIO) vindo do upload
                return Document(file)
        except Exception as e:
            st.error(f"Erro ao carregar documento: {str(e)}")
            return None

    def validate_template(self, doc):
        """Valida o template do documento, retornando lista de problemas."""
        problems = []
        try:
            # Exemplo de validação adicional: capturar placeholders e ver se
            # eles estão bem formados (#ALGO#).
            # Isso complementa a verificação de contagem de '#'.
            # Abaixo mantemos as antigas checagens também:
            for paragraph in doc.paragraphs:
                # Verifica formatação complexa
                if len(paragraph.runs) > 50:
                    problems.append("Parágrafo com formatação muito complexa")

                # Verifica marcadores '#' abertos/fechados
                text = paragraph.text
                if text.count('#') % 2 != 0:
                    problems.append("Marcadores # não estão fechados corretamente")

            # Verifica tabelas complexas (exemplo)
            if len(doc.tables) > 10:
                problems.append("Documento possui muitas tabelas")

        except Exception as e:
            problems.append(f"Erro na validação: {str(e)}")
        return problems

    def unify_and_replace(self, paragraph, old_text, new_text):
        """
        Unifica todos os runs de um parágrafo para texto puro,
        executa a substituição e, em seguida, reatribui o texto em um único run.
        Isso simplifica a preservação de formatação - você perderá formatação
        diferenciada dos runs originais, mas mantém o texto do parágrafo.
        """
        full_text = "".join(run.text for run in paragraph.runs)
        if old_text in full_text:
            new_full_text = full_text.replace(old_text, new_text)
            # Limpa todos os runs
            for run in paragraph.runs:
                run.text = ""
            # Atribui tudo ao primeiro run
            paragraph.runs[0].text = new_full_text

    def replace_text_keeping_format(self, doc, old_text, new_text):
        """
        Substitui o texto em todo o documento, unificando runs por parágrafo.
        Pode causar perda de formatação em runs (negrito, itálico etc.),
        mas evita problemas onde o 'old_text' está quebrado em vários runs.
        """
        for paragraph in doc.paragraphs:
            self.unify_and_replace(paragraph, old_text, new_text)

        # Se quiser substituir também em tabelas:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.unify_and_replace(paragraph, old_text, new_text)

    def auto_backup(self, doc, filename):
        """Cria backup automático do documento"""
        backup_path = Path("backups")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc.save(backup_path / f"{filename}_{timestamp}.docx")
        logging.info(f"Backup criado: {filename}_{timestamp}.docx")

    def scan_for_placeholders(self, doc):
        """
        Retorna um conjunto (set) de placeholders encontrados no doc.
        Ex: #NOME_CLIENTE#, #ENDERECO#, etc.
        """
        placeholders = set()
        pattern = r'#(\w+)#'  # Captura sequências de a-zA-Z0-9 e underscore
        # Parágrafos
        for paragraph in doc.paragraphs:
            for match in re.findall(pattern, paragraph.text):
                placeholders.add(match)
        # Tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for match in re.findall(pattern, paragraph.text):
                            placeholders.add(match)
        return placeholders

    def compare_variables_and_placeholders(self, doc):
        """
        Compara as variáveis definidas em st.session_state.variables com
        os placeholders detectados no doc. Retorna duas listas:
         - missing_in_doc: variáveis definidas que não aparecem no documento
         - missing_in_variables: placeholders no documento que não têm variável definida
        """
        doc_placeholders = self.scan_for_placeholders(doc)
        defined_vars = set(st.session_state.variables.keys())

        missing_in_doc = defined_vars - doc_placeholders
        missing_in_variables = doc_placeholders - defined_vars

        return missing_in_doc, missing_in_variables

    def main(self):
        st.title("🔖 Gerador de Contratos - Versão Aprimorada")

        tabs = st.tabs(["📝 Cadastrar Novo Modelo", "✨ Preencher Contrato"])

        with tabs[0]:
            self.template_tab()

        with tabs[1]:
            self.fill_tab()

    def template_tab(self):
        st.subheader("📝 Cadastrar Novo Modelo")

        # Nome do modelo
        template_name = st.text_input("Nome do Modelo")

        # Upload do arquivo
        uploaded_file = st.file_uploader(
            "Selecione o arquivo .docx",
            type=['docx'],
            help="Apenas arquivos .docx são aceitos"
        )

        if uploaded_file:
            try:
                # Carrega o documento e guarda em memória
                doc = self.load_document(uploaded_file)
                if doc:
                    # Salva conteúdo textual para exibir
                    st.session_state.current_content = '\n'.join(
                        para.text for para in doc.paragraphs
                    )
                    # Salva binário para edição posterior
                    st.session_state.doc_data = uploaded_file.read()
                    st.session_state.original_file = uploaded_file

                    text_area = st.text_area(
                        "Texto do Contrato (exibição simples, sem formatação)",
                        value=st.session_state.current_content,
                        height=300
                    )

                    # Adicionar variável
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        selected_text = st.text_input("Trecho exato do texto para substituir")
                    with col2:
                        var_name = st.text_input("Nome da variável (ex: NOME)").upper()

                    if st.button("➕ Adicionar Variável"):
                        if selected_text and var_name:
                            self.add_variable(selected_text, var_name)

                    st.subheader("📋 Variáveis Adicionadas")
                    if st.session_state.variables:
                        for var, text in st.session_state.variables.items():
                            st.write(f"- **#{var}#** → `{text}`")
                    else:
                        st.write("Nenhuma variável adicionada até o momento.")

                    if st.button("💾 Salvar Modelo"):
                        self.save_template(template_name)
            except Exception as e:
                st.error(f"❌ Erro ao carregar arquivo: {str(e)}")
                logging.error(f"Erro ao carregar arquivo: {str(e)}")

    def fill_tab(self):
        st.subheader("✨ Preencher Contrato")

        # Carregar templates disponíveis
        template_path = Path("templates")
        templates = [f.stem for f in template_path.glob("*.json")]

        if not templates:
            st.warning("⚠️ Nenhum modelo cadastrado ainda!")
            return

        # Seleção do modelo
        selected_template = st.selectbox(
            "Selecione o Modelo",
            options=templates,
            help="Escolha o modelo de contrato"
        )

        if selected_template:
            try:
                with open(template_path / f"{selected_template}.json", "r", encoding="utf-8") as f:
                    template = json.load(f)

                # Criar campos para cada variável
                values = {}
                st.write("Preencha os valores:")
                for var in template["variables"]:
                    values[var] = st.text_input(f"📝 {var}", value="")

                if st.button("🔄 Gerar Contrato"):
                    try:
                        # Carrega o docx correspondente ao template
                        doc_file = template_path / f"{selected_template}.docx"
                        doc = self.load_document(doc_file)
                        if doc:
                            # Substituir variáveis
                            for var_name, value in values.items():
                                if not value:
                                    st.warning(f"⚠️ Campo {var_name} está vazio!")
                                    return
                                placeholder = f"#{var_name}#"
                                self.replace_text_keeping_format(doc, placeholder, value)

                            # Criar backup antes de gerar
                            self.auto_backup(doc, f"{selected_template}_preenchido")

                            # Salvar em memória
                            doc_io = io.BytesIO()
                            doc.save(doc_io)
                            doc_io.seek(0)

                            st.success("✅ Contrato gerado com sucesso!")
                            st.download_button(
                                label="📥 Download Contrato",
                                data=doc_io,
                                file_name=f"{selected_template}_preenchido_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                help="Clique para baixar o contrato gerado"
                            )

                    except Exception as e:
                        st.error(f"❌ Erro ao gerar contrato: {str(e)}")
                        logging.error(f"Erro ao gerar contrato {selected_template}: {str(e)}")

            except Exception as e:
                st.error(f"❌ Erro ao carregar template: {str(e)}")
                logging.error(f"Erro ao carregar template {selected_template}: {str(e)}")

    def add_variable(self, selected_text, var_name):
        """Adiciona uma nova variável e substitui o texto selecionado por #VAR_NAME# no doc em memória."""
        if not var_name:
            st.warning("⚠️ Nome da variável não pode ser vazio!")
            return

        if var_name in st.session_state.variables:
            st.warning("⚠️ Variável já existe!")
            return

        # Registra no dicionário de variáveis
        st.session_state.variables[var_name] = selected_text

        # Se tivermos um doc em memória, tentamos substituir no próprio doc
        if st.session_state.doc_data:
            try:
                doc_bin = io.BytesIO(st.session_state.doc_data)
                doc = Document(doc_bin)

                # Faz a substituição
                self.replace_text_keeping_format(doc, selected_text, f"#{var_name}#")

                # Atualiza o texto atual
                out_bin = io.BytesIO()
                doc.save(out_bin)
                st.session_state.doc_data = out_bin.getvalue()

                # Atualiza a exibição de texto (para feedback ao usuário)
                updated_doc = self.load_document(io.BytesIO(st.session_state.doc_data))
                if updated_doc:
                    st.session_state.current_content = '\n'.join(
                        para.text for para in updated_doc.paragraphs
                    )

                st.success(f"✅ Variável #{var_name}# adicionada e substituída com sucesso!")
                logging.info(f"Variável adicionada: {var_name}")

            except Exception as e:
                st.error(f"Erro ao atualizar documento em memória: {e}")
                logging.error(f"Erro ao atualizar doc em memória: {e}")

    def save_template(self, name):
        """Salva o template com todas as validações."""
        if not name:
            st.warning("⚠️ Digite um nome para o modelo!")
            return

        if not st.session_state.doc_data:
            st.warning("⚠️ Carregue um arquivo primeiro!")
            return

        try:
            doc = self.load_document(io.BytesIO(st.session_state.doc_data))
            if not doc:
                st.error("O documento não pôde ser carregado.")
                return

            # Valida o template
            problems = self.validate_template(doc)
            if problems:
                st.warning("⚠️ Problemas encontrados no template:")
                for p in problems:
                    st.write(f"- {p}")
                if not st.button("Continuar mesmo assim"):
                    return

            # Compara placeholders no doc vs. variáveis definidas
            missing_in_doc, missing_in_variables = self.compare_variables_and_placeholders(doc)
            if missing_in_doc or missing_in_variables:
                st.warning("🚧 Discrepâncias encontradas:")
                if missing_in_doc:
                    st.write("Variáveis definidas, mas ausentes no documento:")
                    for v in missing_in_doc:
                        st.write(f"- {v}")
                if missing_in_variables:
                    st.write("Placeholders no documento sem variáveis definidas:")
                    for v in missing_in_variables:
                        st.write(f"- {v}")
                if not st.button("Salvar mesmo assim"):
                    return

            # Cria backup antes de salvar
            self.auto_backup(doc, name)

            template_path = Path("templates")
            # Dados do template
            data = {
                "variables": list(st.session_state.variables.keys()),
                "last_modified": datetime.now().isoformat(),
                "version": "1.1"
            }

            # Salva metadata
            with open(template_path / f"{name}.json", "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)

            # Salva documento final
            doc.save(template_path / f"{name}.docx")

            st.success("✅ Modelo salvo com sucesso!")
            st.session_state.variables = {}
            st.session_state.current_content = ""
            st.session_state.original_file = None
            st.session_state.doc_data = None

            logging.info(f"Template salvo com sucesso: {name}")

        except Exception as e:
            st.error(f"❌ Erro ao salvar: {str(e)}")
            logging.error(f"Erro ao salvar template {name}: {str(e)}")

def run_app():
    app = GeradorContratosStreamlit()
    app.main()

if __name__ == "__main__":
    run_app()
